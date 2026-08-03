# -*- coding: utf-8 -*-
"""
2026-2학기 외국인 유학생 초대장(LETTER OF INVITATION) — 편집용 Word(.docx) 템플릿 생성
- 기존 합격통지서/비자발급요청서와 동일한 레터헤드(로고)·직인 디자인
- 노란 형광펜 부분만 지원자 정보로 수정해서 개별 발급
- 국문 1페이지 + 영문 1페이지
실행:  python 초대장_docx생성.py [출력폴더]
"""
import sys, os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO = os.path.join(BASE, '로고_국영문_full.png')
SEAL = os.path.join(BASE, '국제협력처장직인.png')

OUT_DIR = sys.argv[1] if len(sys.argv) > 1 else r'C:\Users\user\Desktop\★국제협력처\1. 학부\2026-2학기\초대장'
OUT = os.path.join(OUT_DIR, '초대장_템플릿.docx')

GREEN      = RGBColor(0x1B, 0x6B, 0x3A)
DARKGREEN  = RGBColor(0x12, 0x33, 0x1F)
KEYGREEN   = RGBColor(0x2C, 0x5A, 0x3C)
GRAYGREEN  = RGBColor(0x7D, 0x94, 0x7F)
ADDRGRAY   = RGBColor(0x5A, 0x6B, 0x60)
SIGNGRAY   = RGBColor(0x33, 0x46, 0x3A)
ORGGREEN   = RGBColor(0x20, 0x40, 0x2B)
BLACK      = RGBColor(0x1C, 0x1C, 0x1C)
CELL_BORDER = 'CFE2D5'
KEY_FILL    = 'EAF4EC'

doc = Document()

# ---------- 기본 폰트/페이지 ----------
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
pf = style.paragraph_format
pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.15

sec = doc.sections[0]
sec.page_width = Mm(210); sec.page_height = Mm(297)
sec.left_margin = Mm(19); sec.right_margin = Mm(19)
sec.top_margin = Mm(16); sec.bottom_margin = Mm(14)

CONTENT_MM = 210 - 19*2   # 172mm

# ---------- 헬퍼 ----------
def add_par(container, runs, size=9.5, align=None, before=0, after=0,
            line=1.15, cell=False):
    """runs: list of (text, opts) / opts: b, color, hl, size, spacing(pt)"""
    if cell:
        p = container.paragraphs[0] if not container.paragraphs[0].runs else container.add_paragraph()
    else:
        p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, opts in runs:
        r = p.add_run(text)
        r.font.name = '맑은 고딕'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        r.font.size = Pt(opts.get('size', size))
        r.font.bold = opts.get('b', False)
        r.font.color.rgb = opts.get('color', BLACK)
        if opts.get('hl'):
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if opts.get('spacing'):
            spc = OxmlElement('w:spacing')
            spc.set(qn('w:val'), str(int(opts['spacing'] * 20)))
            r._element.rPr.append(spc)
    return p

def set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_table_borders(table, color, sz=6):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def no_borders_except_bottom(table, color, sz=16):
    tbl = table._tbl
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'none')
        borders.append(el)
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
    borders.append(b)
    tbl.tblPr.append(borders)

def fix_layout(table, widths_mm):
    table.autofit = False
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    table._tbl.tblPr.append(layout)
    for row in table.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)

def cell_margins(table, top=1.6, bottom=1.6, left=2.8, right=2.8):
    tblPr = table._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for name, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement('w:' + name)
        el.set(qn('w:w'), str(int(val * 56.7)))  # mm -> twip
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)

def add_seal(par, size_pt, x_pt, y_pt, uid):
    """서명 문단 기준(paragraph)으로 직인을 '글 뒤' 플로팅 배치"""
    if not os.path.exists(SEAL):
        return
    rId, _ = par.part.get_or_add_image(SEAL)
    cx = cy = int(size_pt * 12700)
    x = int(x_pt * 12700); y = int(y_pt * 12700)
    xml = (
        '<w:drawing ' + nsdecls('w', 'wp', 'a', 'pic', 'r') + '>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        f'relativeHeight="{251658240+uid}" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="margin"><wp:posOffset>{x}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{y}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
        f'<wp:docPr id="{900+uid}" name="seal{uid}"/><wp:cNvGraphicFramePr/>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic><pic:nvPicPr><pic:cNvPr id="{900+uid}" name="seal{uid}"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
        '</a:graphicData></a:graphic></wp:anchor></w:drawing>'
    )
    run = par.add_run()
    run._element.append(parse_xml(xml))

def letterhead(addr):
    t = doc.add_table(rows=1, cols=2)
    no_borders_except_bottom(t, '1B6B3A', sz=16)
    fix_layout(t, [22, CONTENT_MM - 22])
    cell_margins(t, top=0, bottom=1.8, left=0, right=0)
    c0, c1 = t.rows[0].cells
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0); p0.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO):
        p0.add_run().add_picture(LOGO, width=Mm(17))
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_par(c1, [('광주대학교 국제협력처', dict(b=True, color=GREEN, spacing=0.5))], size=12, cell=True)
    add_par(c1, [('Office of International Affairs, Gwangju University', dict(b=True, color=ORGGREEN))], size=9.5, cell=True)
    add_par(c1, [(addr, dict(color=ADDRGRAY))], size=7.5, cell=True, line=1.4)

def info_table(rows, widths):
    t = doc.add_table(rows=len(rows), cols=4)
    set_table_borders(t, CELL_BORDER, sz=6)
    fix_layout(t, widths)
    cell_margins(t)
    for ri, row in enumerate(rows):
        for ci, spec in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if spec is None:
                continue
            kind, runs = spec
            if kind == 'k':
                set_cell_bg(cell, KEY_FILL)
                add_par(cell, runs, size=9, cell=True)
            else:
                add_par(cell, runs, size=9.2, cell=True)
    return t

def K(text):   # key cell
    return ('k', [(text, dict(b=True, color=KEYGREEN))])

def V(runs):   # value cell
    return ('v', runs)

def sign_block(date_txt, uid, before=20):
    add_par(doc, [(date_txt, dict(hl=True, color=SIGNGRAY))], size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER, before=before, after=0)
    p = add_par(doc, [('광주대학교 국제협력처장', dict(b=True, color=DARKGREEN, spacing=4))],
                size=19, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=0)
    # '장' 글자 위에 직인 (content 폭 487.6pt 기준 경험값)
    add_seal(p, 42, 347, -5, uid)
    add_par(doc, [('Director of International Affairs, Gwangju University',
                   dict(b=True, color=SIGNGRAY))], size=11.5,
            align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=0)

def hl(t):  return (t, dict(hl=True))
def txt(t): return (t, dict())
def bold(t): return (t, dict(b=True))

# ============================================================
# PAGE 1 — 국문
# ============================================================
letterhead('(61743) 전남광주통합특별시 남구 효덕로 277 · Tel. +82-62-670-2288 · Fax. +82-62-670-2733 · jhjung@gwangju.ac.kr')

add_par(doc, [('초  대  장', dict(b=True, color=DARKGREEN, spacing=9))], size=22,
        align=WD_ALIGN_PARAGRAPH.CENTER, before=20, after=2)
add_par(doc, [('LETTER OF INVITATION', dict(b=True, color=GRAYGREEN, spacing=2))], size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

add_par(doc, [txt('주 '), hl('베트남'), txt(' 대한민국대사관 영사과 귀중')],
        size=10, before=12, after=0)
p = doc.paragraphs[-1]
for r in p.runs: r.font.bold = True; r.font.color.rgb = ORGGREEN

add_par(doc, [txt('광주대학교는 아래 학생이 '), bold('2026학년도 2학기 외국인 특별전형'),
              txt('에 최종 합격하여 본교에서 정규 학업을 수행할 예정임을 확인하며, 학업 수행을 위하여 아래와 같이 대한민국으로 정식 초청합니다.')],
        size=9.7, before=9, after=0, line=1.5)

add_par(doc, [('■ 피초청인 (초청 대상 학생)', dict(b=True, color=GREEN))], size=9.7, before=11, after=3)
info_table([
    [K('성명'),   V([hl('NGUYEN VAN A'), txt(' ('), hl('구엔 반 아'), txt(')')]), K('생년월일'), V([hl('2005. 1. 1.')])],
    [K('국적'),   V([hl('베트남')]),                                             K('여권번호'), V([hl('C1234567')])],
    [K('지원학과'), V([hl('경영학과')]),                                          K('과정 · 구분'), V([txt('학사과정 · '), hl('신입학')])],
    [K('입학학기'), V([txt('2026학년도 2학기 (2026. 9.)')]),                       K('초청기간'), V([hl('2026. 9. 1. ~ 2030. 8. 31.')])],
], [26, 60, 26, 60])

add_par(doc, [('■ 초청인 (초청기관)', dict(b=True, color=GREEN))], size=9.7, before=11, after=3)
info_table([
    [K('초청기관'), V([txt('광주대학교 (Gwangju University)')]), K('초청인'), V([txt('국제협력처장')])],
    [K('주소'), V([txt('(61743) 전남광주통합특별시 남구 효덕로 277')]), K('연락처'), V([txt('Tel. +82-62-670-2288 · jhjung@gwangju.ac.kr')])],
], [26, 60, 26, 60])

add_par(doc, [txt('위 학생의 입국 목적은 본교에서의 정규 학업 수행에 한정되며, 본교는 학생의 입국 후 학업·출결 및 체류자격 준수 여부를 성실히 관리·감독할 것을 확인합니다. 위 학생에 대한 사증 발급을 긍정적으로 검토하여 주시기를 정중히 요청드립니다.')],
        size=9.7, before=10, after=0, line=1.5)

sign_block('2026. 7. 27.', 1)

# ---------- 페이지 나눔 ----------
doc.add_page_break()

# ============================================================
# PAGE 2 — 영문
# ============================================================
letterhead('277, Hyodeok-ro, Nam-gu, Gwangju 61743, Rep. of Korea · Tel. +82-62-670-2288 · jhjung@gwangju.ac.kr')

add_par(doc, [('LETTER OF INVITATION', dict(b=True, color=DARKGREEN, spacing=2))], size=20,
        align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=2)
add_par(doc, [('초  대  장', dict(b=True, color=GRAYGREEN, spacing=2))], size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

add_par(doc, [('Consular Section, Embassy of the Republic of Korea in ', dict(b=True, color=ORGGREEN)),
              ('Vietnam', dict(b=True, color=ORGGREEN, hl=True))],
        size=10, before=12, after=0)

add_par(doc, [txt('On behalf of Gwangju University, Republic of Korea, we hereby confirm that the student listed below has been officially admitted through the '),
              bold('Special Admission for International Students, Fall Semester 2026'),
              txt(', and formally invite the student to the Republic of Korea for full-time academic study at our university.')],
        size=9.7, before=9, after=0, line=1.5)

add_par(doc, [('■ Invitee (Student)', dict(b=True, color=GREEN))], size=9.7, before=11, after=3)
info_table([
    [K('Name'), V([hl('NGUYEN VAN A')]), K('Date of Birth'), V([hl('Jan. 1, 2005')])],
    [K('Nationality'), V([hl('Vietnam')]), K('Passport No.'), V([hl('C1234567')])],
    [K('Department'), V([hl('Business Administration')]), K('Program'), V([txt("Bachelor's · "), hl('Freshman')])],
    [K('Semester'), V([txt('Fall Semester 2026 (Sep. 2026)')]), K('Period'), V([hl('Sep. 1, 2026 ~ Aug. 31, 2030')])],
], [28, 58, 28, 58])

add_par(doc, [('■ Inviter (Host Institution)', dict(b=True, color=GREEN))], size=9.7, before=11, after=3)
info_table([
    [K('Institution'), V([txt('Gwangju University')]), K('Inviter'), V([txt('Director of International Affairs')])],
    [K('Address'), V([txt('277, Hyodeok-ro, Nam-gu, Gwangju 61743, Rep. of Korea')]), K('Contact'), V([txt('Tel. +82-62-670-2288 · jhjung@gwangju.ac.kr')])],
], [28, 58, 28, 58])

add_par(doc, [txt("We confirm that the sole purpose of the student's entry into the Republic of Korea is to pursue full-time academic study at Gwangju University. The university will faithfully supervise the student's academic performance, attendance, and compliance with visa regulations throughout the period of study. We respectfully request your favorable consideration for the issuance of a visa for the above-mentioned student.")],
        size=9.7, before=10, after=0, line=1.5)

sign_block('July 27, 2026', 2, before=14)

# ---------- 저장 ----------
os.makedirs(OUT_DIR, exist_ok=True)
doc.save(OUT)
print('완료:', OUT)
